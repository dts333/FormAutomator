#!/usr/bin/env python3
# -*- coding: utf-8 -*-
"""
Created on Mon Jun 24 11:53:02 2019

@author: DannySwift
"""

import os
import pandas as pd
import tkinter as tk
from docx import Document
from tkinter import filedialog
from tkinter import *


class gui:
    def __init__(self):
        self.input = None
        self.template = None
        self.output = None
        
        self.top = tk.Tk()
        self.top.title("Form Automator")
        self.input_button = tk.Button(self.top, text='Set CSV', width=25, 
                                      command=self.set_csv)
        self.input_button.grid(row=0, column=0)
        self.template_button = tk.Button(self.top, text='Set Template', 
                                         width=25, command = self.set_template)
        self.template_button.grid(row=1)
        self.output_button = tk.Button(self.top, text='Set Output', width = 25,
                                       command=self.set_output)
        self.output_button.grid(row=2, column=0)
        self.file_name_field = tk.Entry(self.top)
        self.file_name_field.grid(row=3)
        self.run_button = tk.Button(self.top, text='Run', 
                                    width=25, command=self.run)
        self.run_button.grid(row=4)
        #self.add_button = tk.Button(self.top, text='Add Tag', 
        #                            width=25, command=self.add_tag)
        #self.add_button.grid(row=2)
        #self.tags = []
        #self.values = []
        #self.rows = 3
        self.top.mainloop()
    
    
#    def add_tag(self):
#        tk.Label(self.top, text='Tag').grid(row=self.rows)
#        tag = tk.Entry(self.top)
#        tag.grid(row=self.rows, column=1)
#        self.tags.append(tag)
#        tk.Label(self.top, text='Value').grid(row=self.rows, column=2)
#        value = tk.Entry(self.top)
#        value.grid(row=self.rows, column=3)
#        self.values.append(value)
#        self.rows += 1
#        self.top.mainloop()
    
    
    def set_csv(self):
        self.input = filedialog.askopenfilename()
        
    
    def set_template(self):
        self.template = filedialog.askopenfilename()
        
    
    def set_output(self):
        self.output = filedialog.askdirectory()
        
    
    def run(self):
        csv_to_docs(self.input, self.file_name_field.get(), 
                    self.template, self.output)


def format_doc_from_file(file, tags, values):
    doc = Document(file)
    for index in range(len(tags)):
        tag = tags[index]
        value = values[index]
        for p in doc.paragraphs:
            if tag in p.text:
                inline = p.runs
                for i in inline:
                    if tag in i.text:
                        i.text = i.text.replace(tag, value)
    
    return(doc)


def format_doc(doc, tags, values):
    for index in range(len(tags)):
        tag = tags[index]
        value = values[index]
        for p in doc.paragraphs:
            if tag in p.text:
                inline = p.runs
                for i in inline:
                    if tag in i.text:
                        i.text = i.text.replace(tag, value)
    
    return(doc)
 
 
def csv_to_docs(csv, title, template_file, output_dir):
    df = pd.read_csv(csv)
    tags = df.columns
    doc = Document(template_file)
    for _, row in df.iterrows():
        new_doc = format_doc(doc, tags, row.values)
        new_doc.save(output_dir + os.sep + row[title] + '.docx')